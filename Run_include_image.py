import os
from PIL import Image
import openai
from crewai import Agent, Task, Crew
from crewai_tools import (
    DirectoryReadTool,
    FileReadTool,
    SerperDevTool,
    WebsiteSearchTool
)
from docx import Document
from transformers import CLIPProcessor, CLIPModel
import torch

# OpenAI API 
#openai.api_key = 'SCRET_KEY'
#os.environ['OPENAI_API_KEY'] = openai.api_key

openai.api_key = os.getenv('OPENAI_API_KEY')

# Serper API 
SERPER_API_KEY = os.getenv('SERPER_API_KEY')

# Load the CLIP model and processor
clip_model = CLIPModel.from_pretrained("openai/clip-vit-base-patch32")
clip_processor = CLIPProcessor.from_pretrained("openai/clip-vit-base-patch32")

# Function to process an image and generate a description using CLIP
def process_image(image_path):
    image = Image.open(image_path)

    # Define text prompts related to the brand
    text_inputs = [
        "a photo of the brand logo", 
        "a product image from the brand", 
        "an advertisement image from the brand", 
        "a packaging design of the brand", 
        "a promotional material from the brand"
    ]

    # Preprocess image and text inputs
    inputs = clip_processor(text=text_inputs, images=image, return_tensors="pt", padding=True)
    outputs = clip_model(**inputs)
    logits_per_image = outputs.logits_per_image
    
    # Get the best matching description
    best_match_idx = logits_per_image.argmax().item()
    best_description = text_inputs[best_match_idx]
    
    return best_description

# Function to create an SEO agent
def seo_agent():
    goal = "Develop an effective SEO strategy to maximize the brand's online visibility."
    backstory = (
        "The brand is preparing for a new product launch and needs a strong online presence. "
        "SEO is critical to driving organic traffic and ensuring that potential customers discover the brand."
    )
    return Agent(
        role="SEO Researcher",
        goal=goal,
        backstory=backstory,
        tools=[
            SerperDevTool(api_key=SERPER_API_KEY, search_limit=10),
            WebsiteSearchTool()
        ]
    )

# Function to create a market analysis agent
def market_analysis_agent(target_market):
    goal = f"Analyze the market trends, competitive landscape, and consumer behavior in the {target_market} to provide insights that inform the brand's market entry strategy."
    backstory = (
        f"The brand is planning to enter the {target_market} market. To ensure a successful entry, it's crucial to understand the current market trends, identify key competitors, "
        "and assess consumer behavior within this market. This analysis will help the brand to identify opportunities, challenges, and the best strategies for positioning itself in the market."
    )
    return Agent(
        role="Market Analyst",
        goal=goal,
        backstory=backstory,
        tools=[
            DirectoryReadTool(),
            FileReadTool(),
            SerperDevTool(api_key=SERPER_API_KEY, search_limit=10),
            WebsiteSearchTool()
        ]
    )

# Function to create a blog post generation agent
def blog_post_agent():
    goal = "Generate a compelling blog post that communicates the brand's value proposition and engages readers."
    backstory = (
        "The brand aims to attract a new audience segment through high-quality content marketing. "
        "The blog post should be engaging, informative, and optimized for SEO to drive organic traffic."
    )
    return Agent(
        role="Content Writer",
        goal=goal,
        backstory=backstory,
        tools=[
            SerperDevTool(api_key=SERPER_API_KEY)
        ]
    )

# Function to generate a blog post based on research findings
def generate_blog_post(brand_description, image_description, seo_keywords, seo_strategy, market_trends, competitor_analysis):
    context = (
        f"{brand_description}\n"
        f"{image_description}\n"
        f"SEO Keywords: {seo_keywords}\n"
        f"SEO Strategy: {seo_strategy}\n"
        f"Market Trends: {market_trends}\n"
        f"Competitor Analysis: {competitor_analysis}"
    )
    
    response = openai.chat.completions.create(
        model="gpt-4.0",
        messages=[
            {"role": "system", "content": "You are an SEO expert and a market analyst."},
            {"role": "user", "content": context}
        ],
        max_tokens=2000
    )
    
    post_content = response['choices'][0]['message']['content']
    
    return post_content

# Function to save the generated content to a DOCX file
def save_to_docx(post_content, filename='blog_post.docx'):
    doc = Document()
    doc.add_heading('Automated Blog Post', 0)
    doc.add_paragraph(post_content)
    doc.save(filename)
    print(f"File saved as {filename}")

# Main function
def main(brand_description, target_market, image_folder_path):
    # Process all image files in the folder
    all_image_descriptions = []
    
    for image_file in os.listdir(image_folder_path):
        if image_file.lower().endswith(('png', 'jpg', 'jpeg', 'bmp', 'gif')):
            image_path = os.path.join(image_folder_path, image_file)
            image_description = process_image(image_path)
            all_image_descriptions.append(image_description)
    
    # Combine all image descriptions
    combined_image_descriptions = "\n".join(all_image_descriptions)
    
    # Create SEO agent and set up the task
    seo_researcher = seo_agent()
    seo_task = Task(
        agent=seo_researcher,  # Assign the agent here
        objective="Generate SEO keywords and strategy based on brand description.",
        description="Analyze brand's unique selling points and market position to identify the most effective SEO keywords. Formulate a strategy to increase online visibility.",
        expected_output="SEO keyword list, detailed SEO strategy.",
        input_data={"description": brand_description}
    )
    
    # Create market analysis agent and set up the task with target market input
    market_analyst = market_analysis_agent(target_market)
    market_task = Task(
        agent=market_analyst,  # Assign the agent here
        objective=f"Analyze current market trends and competitors in the {target_market} based on brand description.",
        description=f"Research current market trends, identify key competitors, and perform a SWOT analysis to inform the brand's market positioning in the {target_market}.",
        expected_output="Market trend report, competitor SWOT analysis.",
        input_data={"description": brand_description}
    )
    
    # Create blog post generation agent and set up the task
    content_writer = blog_post_agent()
    content_task = Task(
        agent=content_writer,  # Assign the agent here
        objective="Generate a blog post based on the SEO strategy and market analysis findings.",
        description="Utilize SEO keywords and market analysis to write a compelling blog post that engages readers and promotes the brand effectively.",
        expected_output="SEO-optimized blog post draft.",
        input_data={}  
    )
    
    # Set up Crew and execute tasks
    crew = Crew(
        agents=[seo_researcher, market_analyst, content_writer],
        tasks=[seo_task, market_task, content_task],
        verbose=True,
        planning=True,  
        memory=True
    )
    crew.kickoff()  # Just call kickoff without await
    
    # Retrieve task results using .json_dict attribute
    seo_result = seo_task.output.json_dict  
    market_result = market_task.output.json_dict  # 
    post_content_result = content_task.output.json_dict  # 

    # Extract necessary information from task results assuming they are in dict format
    seo_keywords = seo_result.get('keywords', []) if seo_result else []
    seo_strategy = seo_result.get('strategy', 'No strategy generated.') if seo_result else 'No strategy generated.'
    market_trends = market_result.get('market_trends', []) if market_result else []
    competitor_analysis = market_result.get('competitor_analysis', 'No competitor analysis generated.') if market_result else 'No competitor analysis generated.'

    # Generate the final blog post
    post_content = generate_blog_post(
        brand_description, 
        combined_image_descriptions, 
        seo_keywords, 
        seo_strategy, 
        market_trends, 
        competitor_analysis
    )
    
    # Save the result as a DOCX file
    save_to_docx(post_content)

# Execution section
if __name__ == "__main__":
    brand_description = input("Enter brand description: ")
    target_market = input("Enter the target market (e.g., fashion industry, tech industry): ")
    image_folder_path = "./images"  
    
    main(brand_description, target_market, image_folder_path)
